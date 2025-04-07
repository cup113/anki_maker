from html.parser import HTMLParser
from docx.document import Document as DocumentClass
from docx.text.paragraph import Paragraph
from docx.blkcntnr import BlockItemContainer


class HTMLToWordParser(HTMLParser):
    def __init__(self, container: BlockItemContainer | DocumentClass):
        super().__init__()
        self.current_styles: list[dict[str, bool]] = []
        self.container = container
        self.paragraph_stack: list[Paragraph] = []
        self.current_paragraph: Paragraph | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
        style: dict[str, bool] = {}
        if tag == "b" or tag == "strong":
            style["bold"] = True
        elif tag == "i" or tag == "em":
            style["italic"] = True
        elif tag == "u":
            style["underline"] = True
        elif tag == "sub":
            style["subscript"] = True
        elif tag == "p":
            # 创建新段落并推入栈
            if self.current_paragraph is None:
                if isinstance(self.container, BlockItemContainer):
                    new_paragraph = self.container.paragraphs[0]
                else:
                    new_paragraph = self.container.add_paragraph()
            else:
                new_paragraph = self.container.add_paragraph()
            self.paragraph_stack.append(new_paragraph)
            self.current_paragraph = new_paragraph
        self.current_styles.append(style)

    def handle_endtag(self, tag: str):
        if tag in ["b", "i", "u", "strong", "em", "sub"]:
            self.current_styles.pop()
        elif tag == "p":
            # 弹出栈，恢复上一个段落
            if len(self.paragraph_stack) > 1:
                self.paragraph_stack.pop()
                self.current_paragraph = self.paragraph_stack[-1]

    def handle_data(self, data: str):
        if not data:
            return
        current_style: dict[str, bool] = {}
        for style in reversed(self.current_styles):
            current_style.update(style)
        if self.current_paragraph is None:
            if isinstance(self.container, BlockItemContainer):
                self.current_paragraph = self.container.paragraphs[0]
            else:
                self.current_paragraph = self.container.add_paragraph()
        run = self.current_paragraph.add_run(data)
        run.font.bold = current_style.get("bold", False)
        run.font.italic = current_style.get("italic", False)
        run.font.underline = current_style.get("underline", False)
        run.font.subscript = current_style.get("subscript", False)

    @classmethod
    def parse(cls, html: str, container: BlockItemContainer | DocumentClass):
        parser = cls(container)
        parser.feed(html)