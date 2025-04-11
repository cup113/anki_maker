from datetime import datetime

from docx import Document
from docx.document import Document as DocumentClass
from docx.table import Table
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement

from data_models import Chunk, ChunkDocument
from html_parser import HTMLToWordParser


def create_document(chunk_document: ChunkDocument) -> DocumentClass:
    doc = Document()

    # 文档基础设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 页脚配置
    footer = section.footer
    paragraph = footer.add_paragraph()

    # 设置三栏布局的制表位
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(8.3), WD_TAB_ALIGNMENT.CENTER)  # 中间栏
    tab_stops.add_tab_stop(Cm(16.6), WD_TAB_ALIGNMENT.RIGHT)  # 右侧栏

    # 构建页脚内容
    left_run = paragraph.add_run(datetime.now().strftime("%Y.%m.%d"))
    left_run.font.name = "Georgia"
    left_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    paragraph.add_run("\t")  # 第一个制表符

    mid_run = paragraph.add_run("Chunks")
    mid_run.font.name = "Georgia"
    paragraph.add_run("\t")  # 第二个制表符

    right_run = paragraph.add_run()
    right_run.bold = True

    # 插入页码字段（需要XML操作）
    page_fld = OxmlElement("w:fldSimple")
    page_fld.set(qn("w:instr"), "PAGE")
    right_run._r.append(page_fld)
    right_run.add_text(" / ")
    num_fld = OxmlElement("w:fldSimple")
    num_fld.set(qn("w:instr"), "NUMPAGES")
    right_run._r.append(num_fld)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Georgia'  # 英文默认字体
    normal_style.font.size = Pt(11.5)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal_style.paragraph_format.space_before = Pt(1.5)
    normal_style.paragraph_format.space_after = Pt(1.5)
    normal_style.paragraph_format.line_spacing = 1.15

    number_style = doc.styles.add_style('number', WD_STYLE_TYPE.CHARACTER)
    number_style.font.name = 'Cambria'
    number_style.font.size = Pt(11.5)

    # 标题样式设置
    heading_style = doc.styles["Heading 1"]
    heading_style.font.size = Pt(26)
    heading_style.font.bold = False
    heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_style.paragraph_format.space_before = Pt(10)

    heading_style = doc.styles["Heading 2"]
    heading_style.font.size = Pt(18)
    heading_style.paragraph_format.space_before = Pt(10)
    heading_style.paragraph_format.space_after = Pt(4)

    doc.add_heading(level=1).add_run(chunk_document.title).font.name = "Bookman Old Style"

    return doc


def generate_tables(doc: DocumentClass, records: list[Chunk]):
    levels = _group_by_level(records)

    for level, chunks in sorted(levels.items(), key=lambda item: item[0]):
        doc.add_heading(level=2).add_run(level).font.name = "Bookman Old Style"
        table = _create_table(doc, chunks)
        _apply_table_styles(table)
        _fill_table_data(table, chunks)
        doc.add_paragraph()

def generate_footer(doc: DocumentClass, footerHTML: str):
    HTMLToWordParser.parse(footerHTML, doc)

def _group_by_level(chunks: list[Chunk]) -> dict[str, list[Chunk]]:
    grouped: dict[str, list[Chunk]] = {}
    for chunk in chunks:
        grouped.setdefault(chunk.level, []).append(chunk)
    return grouped


def _create_table(doc: DocumentClass, chunks: list[Chunk]) -> "Table":
    table = doc.add_table(rows=len(chunks), cols=3)
    table.autofit = False
    for row in range(len(chunks)):
        for i, width in enumerate([1.0, 8.0, 8.0]):
            table.cell(row, i).width = Cm(width)
            table.cell(row, i).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table

def _apply_table_styles(table: Table):
    tbl = table._element
    tbl_pr = tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    tbl_pr.append(tbl_borders)

    # 边框设置
    _add_border(tbl_borders, "insideH", "single", 4)
    _add_border(tbl_borders, "top", "single", 4)
    _add_border(tbl_borders, "bottom", "single", 4)


def _add_border(tbl_borders: BaseOxmlElement, position: str, style: str, size: int):
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), style)
    border.set(qn("w:sz"), str(size))
    tbl_borders.append(border)


def _fill_table_data(table: Table, chunks: list[Chunk]):
    for i, chunk in enumerate(chunks):
        # 序号列
        table.cell(i, 0).paragraphs[0].add_run(f"{i+1:02d}", style="number")

        # 前置内容
        front_cell = table.cell(i, 1)
        HTMLToWordParser.parse(chunk.get_merged_front(), front_cell)

        # 后置内容
        back_cell = table.cell(i, 2)
        HTMLToWordParser.parse(chunk.get_merged_back(), back_cell)
